# 1. Importando as bibliotecas necessárias para execução do projeto

# criar e gerenciar diretórios/pastas
import os
import shutil

# manipulação e gerenciamento de datas e horários
from datetime import datetime, date

# manipulação de arquivo .docx
from docx import Document

# manipulação de planilhas do excel ou libre calc
from openpyxl import load_workbook


# 2. Caminho dos diretórios
# 2.1 Planilha que será usada como base de dados
DB = 'db.xlsx'

# 2.2 Arquivo de texto contendo os números que serão inseridos
# nos documentos
arq = 'num.txt'

# 2.3 Caminho do diretório principal
MAIN_PATH = f'{os.getcwd()}'

# 2.4 Caminho dos templates
TEMP_PATH = os.path.join(MAIN_PATH, 'templates')

# 2.5 Caminho dos documentos
DOC_PATH = os.path.join(MAIN_PATH, 'doc_emitidos')

# 2.6 Pasta na qual serão incluídos arquivos duplicados
DOC_DOUBLE = os.path.join(MAIN_PATH, 'doc_duplicados')


# 3. Formatando a data
dia = date.today().day
mes = date.today().month
ano = date.today().year
data = date.today()
data_atual = data.strftime('%d%m%y')
hora = datetime.now()
hora_atual = hora.strftime('%H%M')

meses = {
    1: 'janeiro',
    2: 'fevereiro',
    3: 'março',
    4: 'abril',
    5: 'maio',
    6: 'junho',
    7: 'julho',
    8: 'agosto',
    9: 'setembro',
    10: 'outubro',
    11: 'novembro',
    12: 'dezembro'
}


# 4 Criando uma lista com os arquivos do diretório
files = os.listdir(DOC_PATH)
# variável que receberá um valor de um elemento pertecente a lista files
file = f'{ano}' in files

# verificação da existência desse elemento na lista files
if file is True:
    new_folder = file
else:
    new_folder = os.mkdir(f'{DOC_PATH}' + f'\\{ano}')

# 4.1 Criando caminho da nova pasta
YEAR = os.path.join(DOC_PATH, f'{ano}')

# 4.2 Criando um novo diretório para o mês corrente
files = os.listdir(YEAR)
file = f'{ano}{mes}' in files
if file is True:
    new_folder = file
else:
    new_folder = os.mkdir(f'{YEAR}' + f'\\{ano}{mes}')

# 4.3 Criando pasta do mês
MONTH = os.path.join(YEAR, f'{ano}{mes}')

# 4.4 Criando um novo diretório para o dia atual
files = os.listdir(MONTH)

file = f'{dia}' in files

if file is True:
    new_folder = file
else:
    new_folder = os.mkdir(f'{MONTH}' + f'\\{dia}')

DAY = os.path.join(MONTH, f'{dia}')


# 5. Tipos de Documentos
tipo_doc = {
    1: 'Exemplo 1',
    2: 'Exemplo 2',
}


# 6. Lendo os dados da planilha(coluna e celula)
planilha = load_workbook(DB)
plan = planilha['dados']


# 7. Criando uma lista com os dados coletados da planilha
dados = []
dados.clear()
for celula in plan['A'][1:]:
    if not celula.value:
        dados.append(celula.value)


# 8. Obtendo o valor total de itens adicionados à lista
total = len(dados)


# 9. Obtendo uma lista dos documentos presentes no diretório de templates cujo o inicio
# do nome de cada arquivo seja 'doc' e adicionando a um dicionário.
lista = os.listdir(TEMP_PATH)
docs = {}
k = 1
for temp in lista:
    if temp.startswith('doc'):
        docs[k] = temp
    k += 1


# 10. Selecionando o template(documento) que será utilizado para gerar os documentos.
# imprimindo as opções na tela
for k, v in docs.items():
    print(f'[{k:^2}] - {v}')


while True:
    try:
        template = int(input('\nSelecione um documento: '))
    except ValueError:
        print('Digite um número válido!')
    else:
        if template > len(docs) or template <= 0:
            print('Documento inexistente.')
        else:
            break


# 11. Coletando o contexto do template e acrescentando em uma variável
doc2 = Document(f'{os.getcwd()}'+f'\\templates\\{docs.get(template)}')
fullText = []
for linha in doc2.paragraphs:
    fullText.append(linha.text)
texto = '\n'.join(fullText)


# 12. Selecionando o tipo de documento
for k, v in tipo_doc.items():
    print(f'[{k}] - {v}')


while True:
    try:
        tipoDoc = int(input('\nSelecione o tipo de documento: '))
    except ValueError:
        print('Digite um número válido!')
    else:
        if tipoDoc > len(tipo_doc) or tipoDoc <= 0:
            print('Documento inexistente.')
        else:
            break


# 13. Inserindo um número ao documento
while True:
    try:
        num = int(input('Número para o documento: '))
    except ValueError:
        print('Digite um número válido!')
    else:
        break


# 13.1 Informando um assunto para o documento.
assunto = input('Informe o Assunto: ').upper().strip()


# 14. Gerando documentos
for i in range(total):
    # acessando um arquivo modelo para a partir dele criar outros documentos
    doc1 = Document(f'{os.getcwd()}' + '\\modelos\\modelo.docx')
    # adiciona o primeiro parágrafo
    doc1.add_paragraph(f'DOCUMENTO  N. :  {num} / {ano}')
    # adiciona uma linha branco
    doc1.add_paragraph('')
    # adicion o terceiro parágrafo
    doc1.add_paragraph(f'ASSUNTO               :  {assunto}')
    # adiciona uma linha em branco
    doc1.add_paragraph('')
    # adiciona o quinto parágrafo
    doc1.add_paragraph(f'PROTOCOLO/NOME/etc        :  {dados[i]}')
    # formatação dos dois primeiros parágrafos escritos
    doc1.paragraphs[1].runs[0].bold = True
    doc1.paragraphs[3].runs[0].bold = True
    doc1.paragraphs[5].runs[0].bold = True
    # adiciona duas linhas em branco
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    # adicionando o texto
    doc1.add_paragraph(texto)
    # adiciona duas linhas em branco
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    # adicionando data
    doc1.add_paragraph('>>> Ex.: Departamento de Compras.')
    doc1.add_paragraph(f'>>> Local, ao(s) {dia} dia(s) do mês de {meses.get(mes)} de {ano}.')
    # adicionando 5 linhas em branco
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    doc1.add_paragraph('')
    # nome do responsável
    doc1.add_paragraph('>>> Nome do Responsável')
    # formatando a linha
    doc1.paragraphs[18].runs[0].bold = True
    doc1.add_paragraph('>>> Cargo')
    doc1.add_paragraph('Inscrição no Conselho, etc')
    # salvando documento
    # esse formato fica a livre escolha
    doc1.save(f'doc - {dados[i]}{hora_atual}.docx')
    # incrementação do numéro de documento
    num += 1
    # fechando o arquivo
    file.close()


# 15.1 Movendo arquivos
try:
    print('DOCUMENTOS GERADOS COM SUCESSO.')
    docs = os.listdir(MAIN_PATH)
    for doc in docs:
        if doc.startswith('doc -'):
            shutil.move(doc, DAY)
# 15.2 Em caso de documentos duplicados mover arquivos para outra pasta
except ValueError:
    print('DOCUMENTOS JÁ FORAM EMITIDOS.')
    docs = os.listdir(MAIN_PATH)
    for doc in docs:
        if doc.startswith('doc -'):
            shutil.move(doc, DOC_DOUBLE)
